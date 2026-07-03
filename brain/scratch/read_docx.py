import os
import sys
from docx import Document

def docx_to_markdown(docx_path):
    doc = Document(docx_path)
    md_lines = []
    
    # We want to traverse paragraphs and tables in the order they appear.
    # To do that, we can traverse doc.element.body elements.
    for child in doc.element.body:
        if child.tag.endswith('p'):
            # It's a paragraph
            # Find the paragraph object corresponding to this element
            for p in doc.paragraphs:
                if p._element == child:
                    text = p.text.strip()
                    if text:
                        # Simple markdown formatting for headings based on style
                        if p.style and p.style.name and p.style.name.startswith('Heading'):
                            level = p.style.name.replace('Heading', '').strip()
                            try:
                                h_num = int(level)
                                md_lines.append(f"\n{'#' * h_num} {text}\n")
                            except ValueError:
                                md_lines.append(f"\n# {text}\n")
                        else:
                            md_lines.append(text)
                    break
        elif child.tag.endswith('tbl'):
            # It's a table
            for t in doc.tables:
                if t._element == child:
                    md_lines.append("\n| " + " | ".join(cell.text.replace("\n", " ").strip() for cell in t.rows[0].cells) + " |")
                    md_lines.append("| " + " | ".join("---" for _ in t.rows[0].cells) + " |")
                    for row in t.rows[1:]:
                        md_lines.append("| " + " | ".join(cell.text.replace("\n", " ").strip() for cell in row.cells) + " |")
                    md_lines.append("\n")
                    break
    
    # Return full markdown text
    return "\n".join(md_lines)

if __name__ == "__main__":
    import glob
    docx_files = glob.glob(r"Instructions\*.docx")
    os.makedirs(r"brain\scratch", exist_ok=True)
    for f in docx_files:
        basename = os.path.basename(f)
        out_name = f"brain/scratch/{basename.replace('.docx', '.md')}"
        print(f"Converting {f} -> {out_name}")
        try:
            md = docx_to_markdown(f)
            with open(out_name, "w", encoding="utf-8") as out_f:
                out_f.write(md)
        except Exception as e:
            print(f"Error converting {f}: {e}")
